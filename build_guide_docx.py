"""Generate PROJECT_GUIDE.docx, the full build-from-scratch walkthrough.

Kept as a script rather than a hand-edited binary so the guide can be regenerated
whenever the project changes.

    uv run python build_guide_docx.py
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = "PROJECT_GUIDE.docx"

MONO = "Menlo"
CODE_BG = "F4F4F4"
ACCENT = RGBColor(0x1A, 0x53, 0x8C)
MUTED = RGBColor(0x55, 0x55, 0x55)


def shade(paragraph, hex_fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    paragraph._p.get_or_add_pPr().append(el)


def code(doc, text):
    for line in text.strip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Pt(14)
        run = p.add_run(line or " ")
        run.font.name = MONO
        run.font.size = Pt(8.5)
        shade(p, CODE_BG)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def body(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.italic = italic
    return p


def bullets(doc, items):
    """Bullet list. A short lead-in before the first comma is bolded as a label.

    Guarded on length and punctuation so it only fires on real "label, explanation"
    bullets and leaves ordinary comma-containing sentences alone.
    """
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        head, sep, tail = item.partition(", ")
        if sep and len(head) <= 34 and ":" not in head and not head.endswith("."):
            r = p.add_run(head)
            r.bold = True
            r.font.size = Pt(10.5)
            p.add_run(sep + tail).font.size = Pt(10.5)
        else:
            p.add_run(item).font.size = Pt(10.5)


def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.italic = True
    run.font.color.rgb = MUTED


def phase(doc, title):
    doc.add_page_break()
    h = doc.add_heading(title, level=1)
    h.runs[0].font.color.rgb = ACCENT


def sub(doc, title):
    h = doc.add_heading(title, level=2)
    h.runs[0].font.size = Pt(12)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for cell, text in zip(t.rows[0].cells, headers):
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = ""
            cell.paragraphs[0].add_run(str(text)).font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ---------------------------------------------------------------- title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Abaid Automobile Showroom")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = ACCENT

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Building a semantic car recommender with an AI sales agent\n"
        "from an empty folder to a working app"
    )
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED

    doc.add_paragraph()
    body(doc,
         "This guide rebuilds the project end to end. Every phase is a checkpoint: "
         "run the commands, confirm the stated output, then move on. Phases 7 to 10 "
         "are the interesting ones: they are the bugs the first working version "
         "shipped with, and why the fixes take the shape they do.")

    sub(doc, "What you end up with")
    bullets(doc, [
        "A Streamlit app, search grid on the left, chat with a sales agent on the right.",
        "A search engine that treats stated requirements as filters and everything else as semantics.",
        "A LangGraph ReAct agent with four tools, grounded in the inventory.",
        "Clickable result cards that open a full detail modal.",
        "~14,800 listings, embedded and searchable.",
        "40 offline tests.",
    ])

    sub(doc, "Prerequisites")
    bullets(doc, [
        "uv, the Python package manager (astral.sh/uv).",
        "Python 3.12, uv installs it if missing.",
        "An OpenAI API key, used for embeddings (~$0.05 total) and the chat agent.",
        "A Kaggle account, kagglehub prompts for credentials on first download.",
        "Apple Silicon or an NVIDIA GPU, for the vibe-tagging step. CPU works but is very slow.",
    ])

    # ------------------------------------------------------------------ phase 0
    phase(doc, "Phase 0: Environment scaffold")
    body(doc,
         "Open VS Code, then open the parent folder where you keep projects "
         "(e.g. ~/Projects), not the car project itself, since it does not exist yet.")
    body(doc, "Open the integrated terminal (Ctrl+`) and run:")
    code(doc, """
# confirm uv is installed
uv --version

# scaffold the project (creates pyproject.toml, .python-version, main.py, .gitignore)
uv init car-recom-system --python 3.12
cd car-recom-system

# core deps for the pipeline and the app
uv add pandas numpy python-dotenv streamlit \\
       langchain langchain-community langchain-openai \\
       langchain-chroma langchain-text-splitters chromadb
uv add --dev ipykernel jupyter

git init
""")
    body(doc,
         "Once the folder exists, switch VS Code into it properly: File → Open Folder "
         "→ select car-recom-system. This reloads the window with the project as the "
         "workspace root, which matters for interpreter detection, the .venv picker, "
         "and the Jupyter extension later.")
    code(doc, """
printf ".venv/\\n.env\\n__pycache__/\\n*.pyc\\n.DS_Store\\nchroma_db/\\n" >> .gitignore
git add .
git commit -m "chore: initial scaffold via uv init"
""")
    body(doc,
         'In VS Code: Cmd+Shift+P → "Python: Select Interpreter" → pick '
         "car-recom-system/.venv/bin/python so it uses the uv-managed environment.")
    body(doc, "For GitHub, if you have the gh CLI:")
    code(doc, "gh repo create car-recom-system --public --source=. --remote=origin --push")
    body(doc,
         "Otherwise create the empty repo on github.com first, then "
         "git remote add origin <url> && git push -u origin main.")

    # ------------------------------------------------------------------ phase 1
    phase(doc, "Phase 1: Get the data and look at it")
    body(doc, "Run this first, in a notebook, to see what you are working with:")
    code(doc, """
uv add kagglehub
uv run jupyter notebook
""")
    body(doc,
         "That opens Jupyter using the uv-managed env. The uv run prefix makes sure "
         "it is the project's .venv and not some other Python on your system. Create "
         "a notebook called data-exploration.ipynb, select the matching kernel, and run:")
    code(doc, """
import kagglehub
import pandas as pd

path = kagglehub.dataset_download("austinreese/craigslist-carstrucks-data")
print("Path to dataset files:", path)

cars = pd.read_csv(f"{path}/vehicles.csv")
print(cars.shape)
print(cars.columns.tolist())
cars.head()
""")
    note(doc, "Expect (426880, 26). The CSV is about 1.4 GB, so the first download takes a while.")

    sub(doc, "Look before you clean")
    code(doc, """
print(cars["description"].isna().sum(), "missing descriptions")
print(cars["description"].str.split().str.len().describe())

for col in ["type", "fuel", "transmission", "drive", "condition"]:
    print(f"--- {col} ---")
    print(cars[col].value_counts(dropna=False))

# this dataset is notorious for reposted listings across regions
print("duplicate descriptions:", cars["description"].duplicated().sum())
""")
    body(doc, "Three things that shape every later decision:")
    bullets(doc, [
        "Heavy duplication: dealers repost the same ad across regions.",
        "Description length is bimodal: a handful of words at one end, 2000-word dealer boilerplate at the other.",
        "fuel is extremely imbalanced and gas dominates. Remember this; it causes a real bug in Phase 7.",
    ])

    # ------------------------------------------------------------------ phase 2
    phase(doc, "Phase 2: Clean and sample")
    body(doc,
         "426k rows is more than you need and most of it is noise. Filter to listings "
         "that carry real signal:")
    code(doc, """
import re

cars = cars[cars["description"].notna()].copy()
cars = cars.drop_duplicates(subset="description", keep="first")

# cuts near-empty ads AND spam-length boilerplate
word_count = cars["description"].str.split().str.len()
cars = cars[word_count.between(15, 200)]

cars = cars[cars["price"].between(500, 100_000)]
cars = cars[cars["odometer"].between(500, 300_000)]
cars = cars[cars["year"].between(1990, 2022)]

print(cars.shape)   # ~(108390, 26)
""")
    code(doc, """
def clean_description(text):
    text = re.sub(r'http\\S+|www\\.\\S+', ' ', text)                    # URLs
    text = re.sub(r'\\(?\\d{3}\\)?[-.\\s]?\\d{3}[-.\\s]?\\d{4}', ' ', text)  # phone numbers
    text = re.sub(r'[\\u260f\\u2705\\U0001F698\\U0001F4DE]+', ' ', text)  # symbols
    return re.sub(r'\\s+', ' ', text).strip()

cars["description_clean"] = cars["description"].apply(clean_description)
""")

    sub(doc, "The sampling decision that causes a bug later")
    body(doc,
         "108k rows is still more than needed, so sample down. The obvious move is to "
         "stratify by body type so you get an even spread of sedans, SUVs and trucks:")
    code(doc, """
sampled_idx = (
    cars.groupby("type")
    .apply(lambda g: g.sample(n=min(len(g), 1000), random_state=42).index)
)
cars_sampled = cars.loc[np.concatenate(sampled_idx.values)].reset_index(drop=True)
print(cars_sampled["fuel"].value_counts())
""")
    body(doc,
         "This is correct for body type and quietly wrong for everything else. "
         "Stratifying on type says nothing about fuel, so the result is 90.6% gas with "
         "37 electric cars out of 12,429, and exactly one electric sedan. Phase 8 "
         "fixes it. The final pipeline/build_dataset.py does both steps at once, but "
         "it is worth seeing why the correction is needed.")

    # ------------------------------------------------------------------ phase 3
    phase(doc, "Phase 3: Vibe tagging")
    body(doc,
         "People shop by feel, 'something good for the family', 'a fun weekend car'. "
         "Score every listing against six vibes using zero-shot classification, so "
         "that language has something to match.")
    code(doc, """
uv add transformers torch tqdm
""")
    body(doc, "Create pipeline/vibe_tagging.py:")
    code(doc, """
VIBE_LABEL_MAP = {
    "a family-friendly car with practical space": "vibe_family",
    "a sporty high-performance car":              "vibe_sporty",
    "a luxury car with premium features":         "vibe_luxury",
    "a fuel-efficient eco-friendly car":          "vibe_ecofriendly",
    "an off-road or rugged utility vehicle":      "vibe_offroad",
    "a basic reliable daily commuter car":        "vibe_commuter",
}
HYPOTHESIS = "This vehicle listing describes {}."

device = "mps" if torch.backends.mps.is_available() else "cpu"
classifier = pipeline("zero-shot-classification",
                      model="facebook/bart-large-mnli", device=device)
""")
    note(doc,
         "About 7 minutes per 1,000 rows on Apple MPS, roughly 110 minutes for the "
         "full table. Checkpoint every 20 batches and support resuming; you will "
         "interrupt it at least once.")
    body(doc,
         "One subtlety: multi_label=False makes the six scores a softmax across the "
         "labels, so they are competitive shares, not independent confidences. "
         "vibe_family averages 0.31 and is the top label for 37% of rows. That makes "
         "them a weak ranking prior, not a filter, which is why Phase 6 z-scores "
         "them before blending.")

    # ------------------------------------------------------------------ phase 4
    phase(doc, "Phase 4: Embeddings")
    body(doc,
         "Semantic search needs a vector per listing. What you put into that vector "
         "matters more than which model you use.")
    body(doc, "Create pipeline/build_embeddings.py:")
    code(doc, """
def build_embedding_text(row):
    return (
        f"{year} {manufacturer} {model}. "
        f"{type} body, {fuel} fuel, {transmission} transmission, "
        f"{drive} drivetrain, {condition} condition, {mileage} miles. "
        f"{description_clean}"
    )
""")
    body(doc,
         "Lead with the structured facts. The first version of this project embedded "
         "the row id plus the raw Craigslist blurb, which meant the clean fuel, "
         "transmission and type columns were never in the vectors at all, so a query "
         "saying 'automatic' had nothing reliable to match against, and the leading "
         "numeric id was pure noise.")
    code(doc, """
uv run python -m pipeline.build_embeddings
""")
    note(doc,
         "Use text-embedding-3-small: 1536 dimensions, better retrieval than "
         "ada-002 and five times cheaper. The whole table costs about $0.05.")

    sub(doc, "Keep the rows aligned")
    body(doc,
         "Build the array as a plain positional list and save it directly. If you "
         "round-trip through a vector store, do not assume .get() returns insertion "
         "order, and never trust a row-count assert to prove alignment, because it "
         "only checks length. A silent off-by-N here is very hard to debug later.")

    # ------------------------------------------------------------------ phase 5
    phase(doc, "Phase 5: The search engine")
    body(doc, "Create showroom/search_engine.py. Cosine similarity is the core:")
    code(doc, """
query_vec = np.array(embedder.embed_query(semantic_query))
sims = (embeddings @ query_vec) / (
    np.linalg.norm(embeddings, axis=1) * np.linalg.norm(query_vec) + 1e-8
)
""")
    body(doc,
         "Add exact filters for the categorical columns and a price range, then sort. "
         "At this point you have a working recommender, and it is subtly broken in "
         "ways that only show up once you use it properly. That is Phase 7.")

    # ------------------------------------------------------------------ phase 6
    phase(doc, "Phase 6: The Streamlit app")
    code(doc, "uv run streamlit run app.py")
    body(doc, "Structure that matters:")
    bullets(doc, [
        "@st.cache_data for the dataframe and embeddings: they load once, not per interaction.",
        "@st.cache_resource for the embedding client.",
        "st.form so the search only fires on submit, not on every keystroke.",
        "st.session_state for anything that must survive a rerun, Streamlit re-executes the whole script top to bottom on every interaction.",
    ])
    note(doc,
         "use_container_width is deprecated. Use width=\"stretch\" instead.")

    # ------------------------------------------------------------------ phase 7
    phase(doc, "Phase 7: Diagnosing the flaws")
    body(doc,
         "The app works, so try a realistic query. Use one where you state several "
         "requirements at once, the way a real customer would:")
    code(doc, '"family car sedan with good fuel mileage. I want automatic transmission. Electric car."')
    body(doc, "Top 12 results: zero electric, three manual, one sedan. Every stated "
              "requirement ignored. Four separate causes:")

    sub(doc, "1. Nothing enforces constraints from free text")
    body(doc,
         "The query is embedded whole and ranked by cosine similarity. Words like "
         "'automatic' and 'electric' are just tokens averaged into one vector. The "
         "only real filtering lives in the sidebar dropdowns.")

    sub(doc, "2. Multi-clause queries dilute every clause")
    table(doc, ["query", "result"], [
        ['"electric car" alone', "top 8 all electric, similarity 0.84–0.87"],
        ["the full four-clause query", "best electric ranks #38, behind 37 gas cars"],
    ])
    body(doc,
         "The index knows perfectly well which cars are electric. Averaging four "
         "clauses into one vector is what destroys the signal.")

    sub(doc, "3. Similarity has almost no dynamic range")
    body(doc,
         "Across all rows, similarities span 0.691 to 0.836 with a standard deviation "
         "of 0.0185. Every car is 'sort of similar' to any car query, so rank ordering "
         "inside that band is mostly noise: the gap between #1 and #38 is 0.015.")

    sub(doc, "4. The sampling left nothing to find")
    body(doc,
         "Only 37 electric cars out of 12,429, and exactly one electric sedan. Even a "
         "perfect ranker cannot return cars that are not in the table.")

    sub(doc, "Also: the score blend is backwards")
    body(doc,
         "0.7 * similarity + 0.3 * vibe looks like it weights similarity more heavily. "
         "It does not. Similarity spans about 0.10 while vibe spans the full 0–1, so "
         "the 0.3 term drives roughly three times more variance. The weights meant the "
         "opposite of how they read.")

    # ------------------------------------------------------------------ phase 8
    phase(doc, "Phase 8: Fixing it")

    sub(doc, "Extract constraints from the query text")
    body(doc, "Create showroom/query_parser.py. Map phrasings onto the closed set of "
              "values the columns actually contain:")
    code(doc, """
FUEL_PATTERNS = [
    ("electric", r"\\b(?:all[-\\s]?electric|battery[-\\s]?electric|electric|bev|ev)\\b"),
    ("hybrid",   r"\\b(?:plug[-\\s]?in\\s+hybrid|phev|hybrid)\\b"),
    ("diesel",   r"\\b(?:diesel|tdi)\\b"),
    ("gas",      r"\\b(?:gasoline|petrol|gas(?:\\s+powered)?)\\b"),
]
""")
    body(doc,
         "Then strip the matched phrases from the text you embed. The constraints are "
         "already enforced exactly, so spending the semantic ranking on them again is "
         "wasted, and actively harmful, since it is what diluted the query.")

    sub(doc, "Guard the false positives")
    body(doc,
         "'good gas mileage' is a request for efficiency, not a demand for a gasoline "
         "engine. Without a guard, the single commonest way people ask for a "
         "fuel-efficient car silently filters out every hybrid and EV.")
    code(doc, """
GAS_FALSE_POSITIVES = re.compile(
    r"\\b(?:good|great|better|best|excellent|decent|easy|cheap|low|efficient)\\b"
    r"[^.]{0,20}\\bgas\\b"
    r"|\\bgas\\s+(?:mileage|saver|savings|efficient|economy|sipper)\\b",
    re.IGNORECASE,
)
""")

    sub(doc, "Relax honestly when nothing matches")
    body(doc,
         "If no car satisfies every requirement, drop them one at a time, weakest "
         "first, and tell the user what was dropped. Never silently return something "
         "that violates what they asked for.")
    code(doc, 'RELAX_ORDER = ["drive", "type", "transmission", "fuel"]')
    body(doc,
         "Drivetrain is the most incidental ask; fuel is the most deliberate. Someone "
         "who typed 'electric' did not mean 'or gas is fine'.")

    sub(doc, "Fix the blend")
    code(doc, """
def _zscore(series):
    std = series.std()
    if not np.isfinite(std) or std < 1e-9:
        return pd.Series(np.zeros(len(series)), index=series.index)
    return (series - series.mean()) / std

score = 0.7 * _zscore(similarity) + 0.3 * _zscore(vibe)
""")

    sub(doc, "Rebuild the data")
    body(doc,
         "Keep the type-stratified base, then top up the rare fuels the sampling threw "
         "away. See pipeline/build_dataset.py.")
    table(doc, ["", "before", "after"], [
        ["total rows", "12,429", "14,828"],
        ["electric", "37", "284"],
        ["electric sedans", "1", "71"],
        ["hybrid", "175", "1,346"],
    ])
    body(doc, "Then re-embed everything with the structured text from Phase 4.")

    sub(doc, "Result")
    table(doc, ["", "before", "after"], [
        ["electric", "0 / 12", "10 / 10"],
        ["automatic", "9 / 12", "10 / 10"],
        ["sedan", "1 / 12", "10 / 10"],
    ])

    # ------------------------------------------------------------------ phase 9
    phase(doc, "Phase 9: Budget parsing")
    body(doc,
         "Ask for 'cars under 5000 dollars' and you get cars at $6,200 and $7,500. The "
         "parser handles fuel, transmission, body type and drivetrain, but not price.")
    body(doc, "It is worse than merely unfiltered. The digits go into the embedder, and "
              "'5000' matches the Fiat 500 and Ford 500 model names:")
    code(doc, """
before:  $5,999 Mazda 3 | $6,200 Fiat 500 | $750 Cavalier | $3,900 Ford 500
         $2,500 Fiat 500c | $5,480 Fiat 500 | $4,500 Fiat 500 | $7,500 Fiat 500c
         -> 4 of 8 over budget, 5 of 8 are "500"s

after:   $1,850 | $1,995 | $2,500 | $3,200 | $3,495 | $3,595 | $4,499 | $4,990
         -> 0 violations
""")
    body(doc, "Two guards are essential:")
    bullets(doc, [
        '"under 50,000 miles", mileage, not a budget. Reject when a mileage unit follows.',
        '"cars under 2015", a model year, not a budget. Reject bare numbers in 1900–2100 unless a currency cue is present ($, dollars, or a k suffix).',
    ])
    body(doc,
         "Strip the digits from the embedded text unconditionally. This needs care: if "
         "your parser falls back to the original query when stripping leaves too "
         "little, that fallback will put the digits straight back. Budget spans must "
         "never be restored, embed a neutral phrase instead, since the filter has "
         "already done the real work.")
    body(doc,
         "Finally, intersect the parsed budget with the sidebar slider rather than "
         "letting one override the other. Both are real constraints; the tighter bound "
         "wins. And never relax a budget, someone who says 'under $5,000' means it.")

    # ----------------------------------------------------------------- phase 10
    phase(doc, "Phase 10: The sales agent")
    code(doc, "uv add langgraph")
    body(doc, "Create showroom/sales_agent.py with a LangGraph ReAct agent:")
    code(doc, """
from langgraph.prebuilt import create_react_agent

def build_agent(cars, embeddings, embedder, sink):
    tools = build_tools(cars, embeddings, embedder, sink)

    def prompt(state):
        return [SystemMessage(SALES_PROMPT + describe_grid(sink))] + state["messages"]

    return create_react_agent(ChatOpenAI(model=CHAT_MODEL), tools, prompt=prompt)
""")

    sub(doc, "Four tools")
    table(doc, ["tool", "what it does"], [
        ["search_inventory", "searches the lot AND repaints the results grid"],
        ["get_car_details", "one listing in full, by id"],
        ["compare_cars", "2–4 listings side by side"],
        ["inventory_stats", "aggregates over the whole lot; does not touch the grid"],
    ])

    sub(doc, "The sink pattern")
    body(doc,
         "Tools close over a plain dict. When the agent searches, it writes results "
         "there, and Streamlit reads it to repaint the grid, which is how 'show me "
         "something cheaper' updates the screen instead of only being described in "
         "prose. Hold it in st.session_state and mutate it in place, never reassign: "
         "the tools captured that exact object.")

    sub(doc, "Grounding is the whole game")
    bullets(doc, [
        "Only discuss cars a tool returned this conversation. Never invent a listing, price or spec.",
        "Structured fields are reliable; the seller's description is hearsay. Quote it as \"the listing says\".",
        'A field reading "unknown" means the seller never stated it, never fill it in from the model name.',
        "When a requirement was relaxed, results are NOT guaranteed to match it. Do not summarise the group as though they do.",
        "No service history, accident records or real-world range exist in this data. Say so.",
    ])
    note(doc,
         "That fourth rule was added after the agent described a relaxed search as "
         "'diesel 4x4 trucks' when drivetrain had just been dropped and two of six rows "
         "had drive=unknown. It is a subtler version of the original bug, worth "
         "testing for explicitly.")

    sub(doc, "Make failed searches recoverable")
    body(doc,
         "An empty result should tell the agent which requirement is the blocker, so "
         "it can offer a real alternative:")
    code(doc, """
NOTE: nothing matches all of that. dropping fuel=electric would give 1,135 car(s);
dropping body_type=pickup would give 284 car(s). Tell the customer what is
unavailable and offer the closest alternative - do not pretend it matches.
""")

    sub(doc, "Normalise the tool arguments")
    body(doc,
         "The dataset says 'SUV' but 'sedan', and lowercases manufacturers. The model "
         "has no way to guess that, so map whatever it sends onto the real values "
         "case-insensitively, otherwise body_type='suv' silently matches nothing.")

    sub(doc, "Rendering the replies: the dollar-sign trap")
    body(doc,
         "The first version of the chat produced text like this on screen:")
    code(doc, """
The range runs from
3,000 ** to ** 3,000 ** to ** 22,300, with a couple standouts being the 2018
Honda Civic Sport Hatchback at 14,995 ** with ** 28,123 miles ** and the **
""")
    body(doc,
         "It looks like the model is emitting garbage. It is not. Streamlit's markdown "
         "renders $...$ as LaTeX math, so the moment a reply quotes two prices, "
         "everything between the first and second dollar sign is parsed as an "
         "equation, the asterisks survive as literal math symbols and fragments get "
         "duplicated. A salesperson quotes two prices in almost every sentence, so "
         "this fires constantly.")
    body(doc, "Escape the dollar signs before rendering:")
    code(doc, r'''
def md_safe(text):
    """Streamlit treats `$...$` as LaTeX. Escape so prices render as prices."""
    return str(text).replace("$", r"\$")

st.chat_message("assistant").write(md_safe(msg.content))
''')
    note(doc,
         "Apply it to the user's messages too, \"under $5,000 and over $1,000\" "
         "breaks in exactly the same way.")
    body(doc,
         "Separately, tell the agent to drop markdown altogether. A narrow chat "
         "column full of bold runs reads as clutter even when it renders correctly, "
         "and plain prose is what a salesperson actually sounds like:")
    code(doc, """
- Write plain prose. No markdown: no **bold**, no bullet points, no headings, no
  tables. You are speaking to someone across a desk, not writing a document.
- Write prices as plain numbers with a dollar sign and commas, e.g. $14,995.
""")
    body(doc, "The two fixes are independent and you need both, escaping alone still "
              "leaves the clutter, and the prompt alone still trips LaTeX on any "
              "sentence containing two prices.")

    # ----------------------------------------------------------------- phase 11
    phase(doc, "Phase 11: The detail modal")
    body(doc,
         "Cards show a summary. Clicking one should open the full listing. Streamlit "
         "has a native modal that supplies its own close button, so this needs no "
         "custom component:")
    code(doc, """
@st.dialog("Vehicle details", width="large", on_dismiss="rerun")
def show_car_details(car_id):
    row = cars[cars["id"] == car_id].iloc[0]
    st.markdown(f"### {car_title(row)}")
    ...
""")

    sub(doc, "Opening it without it reopening forever")
    body(doc,
         "A button sets the selection; code after the grid opens the dialog. The "
         "detail that matters is popping the state rather than reading it, if you "
         "read it, the modal reopens on the very next rerun after the user closes it, "
         "which is the classic version of this bug.")
    code(doc, """
# inside the card
if st.button("View full details", key=f"card_{row['id']}", width="stretch"):
    st.session_state.selected_car = int(row["id"])

# after the grid, popped, not read
if "selected_car" in st.session_state:
    show_car_details(st.session_state.pop("selected_car"))
""")
    note(doc,
         "Key the buttons by listing id, never by loop index. Index keys collide the "
         "moment a new search reorders the grid and you open the wrong car. Also note "
         "only one dialog may be called per script run.")

    sub(doc, "What goes in it")
    bullets(doc, [
        "Four headline metrics, price, odometer, year, condition, as bordered st.metric tiles.",
        "A two-column specification grid: body type, fuel, transmission, drivetrain, cylinders, size, paint, title status, VIN, posting date.",
        "A horizontal bar chart of the six vibe scores, captioned as model-inferred impressions rather than manufacturer specs.",
        "The seller's description in a scrollable bordered container, captioned as seller-written.",
    ])

    sub(doc, "The photo that wasn't there")
    body(doc,
         "The dataset has an image_url on 100% of rows, so photos look like a free "
         "win. Check before building on it:")
    code(doc, """
# HEAD returns 405, that is "method not allowed", NOT a dead link.
# Re-check with a GET range request before concluding anything:
req = urllib.request.Request(url, headers={"Range": "bytes=0-2047",
                                           "User-Agent": "Mozilla/5.0"})
""")
    body(doc,
         "Result: 0 of 20 images resolve, and 0 of 3 listing pages resolve. Every URL "
         "404s. This is a 2021 snapshot and Craigslist purges images when a listing "
         "expires. Wiring image_url into the grid would have produced a page of "
         "broken-image icons, and the 'open original listing' button would have been "
         "a link that always fails.")
    body(doc,
         "The tempting fix is a stock photo looked up by make, model and year. Do not. "
         "Two reasons, and the second is the real one:")
    bullets(doc, [
        "It would barely work, 4,284 distinct model strings across 14,828 rows, 2,820 of them appearing exactly once, in dirty free text like \"wrangleryj\" and \"veloster 3dr cpe auto\", with 5.8% missing a manufacturer entirely.",
        "It would lie. A studio render of a clean 2004 Prius above \"200,000 miles, condition fair, salvage title\" tells the customer they are looking at that car. Used cars are individual; the photo is where a buyer reads wear and damage. This is the same failure as returning a manual when the customer asked for an automatic, just better dressed.",
    ])
    body(doc,
         "The project shipped with no images at all. That is the honest answer when "
         "the media is gone, and it costs nothing that matters, the modal already "
         "carries every fact the data actually holds.")

    sub(doc, "Removing the sidebar filters")
    body(doc,
         "The original app had eight sidebar filter widgets. Once the query parser "
         "extracts fuel, transmission, body type, drivetrain and budget from the "
         "sentence, and the agent passes the same constraints through its tools, the "
         "dropdowns duplicate both paths. Delete them.")
    note(doc,
         "Watch for the controls in there that are not filters. \"Clear conversation\" "
         "moved next to the chat it clears; the result count became a constant. "
         "Deleting a panel wholesale quietly removes whatever else was living in it.")

    # ----------------------------------------------------------------- phase 12
    phase(doc, "Phase 12: Collapsing to one input")
    body(doc,
         "At this point the app has two ways to ask for a car: a search box and a "
         "chat with Sam. They do the same job, so one of them has to go. Sam stays, "
         "because he can also answer questions the box never could.")

    sub(doc, "You probably do not need a new tool")
    body(doc,
         "The instinct is to add a filtering tool to the LangGraph workflow so the "
         "agent can drive the grid. That work is already done. search_inventory "
         "writes its results to the sink, and Streamlit renders the grid from the "
         "sink, so the agent has been repainting the screen since the moment it was "
         "built. Removing the search box is deletion, not addition:")
    bullets(doc, [
        "The st.form search box and its run_box_search() function, duplicating what search_inventory already does.",
        "The entire from_box mechanism, which existed only so Sam could comment on a box search without re-running it. No box, no need.",
        "The search_cars import in app.py, since the app no longer searches directly at all.",
    ])
    body(doc,
         "The result is one search path instead of two implementations that had to "
         "be kept in step. That whole class of bug disappears by construction.")

    sub(doc, "The gap deletion opens")
    body(doc,
         "One thing does break, and it is easy to miss. The banners that report what "
         "was enforced and what had to be relaxed are fed by sink[\"meta\"], which only "
         "the search box was writing. Delete the box and the app silently loses the "
         "transparency the whole project was built around. Have the tool write it:")
    code(doc, """
explicit = {k: v for k, v in _spec(fuel, transmission, body_type, drive,
                                   manufacturer).items() if v}
found.constraints = {**explicit, **found.constraints}
sink["meta"] = found
""")
    note(doc,
         "Note the merge. search_cars deliberately keeps explicitly-passed filters "
         "out of `constraints`, because it treats them as deliberate choices rather "
         "than inferred ones. The agent passes constraints explicitly, so without "
         "folding them back in the banner under-reports exactly the requirements the "
         "customer stated most plainly.")

    sub(doc, "Say what is on the lot")
    body(doc,
         "With no filter dropdowns left, nothing tells the customer what the "
         "dealership actually stocks. Put the totals in the header:")
    code(doc, """
@st.cache_data
def inventory_summary():
    counts = cars["type"].value_counts()
    vague = [t for t in ("other", "unknown") if t in counts.index]
    named = counts.drop(index=vague)
    return len(cars), pd.concat([named, counts[vague]])

st.markdown(f"**{total_cars:,} vehicles in stock.**")
with st.container(horizontal=True):
    for body_type, count in type_counts.items():
        st.badge(f"{body_type} {count:,}", color="gray")
""")
    note(doc,
         "Sorting purely by count puts \"unknown\" first, because it is the single "
         "largest bucket. That is a poor opening line for a showroom, so real body "
         "types lead and the catch-alls go last.")
    body(doc,
         "Greet the customer with the same number when the page loads, so the "
         "conversation opens with something concrete rather than a blank prompt:")
    code(doc, """
"Welcome to Abaid Automobile Showroom. I'm Sam. We have 14,828 vehicles on the
lot right now. How can I help you today?"
""")

    # ----------------------------------------------------------------- phase 13
    phase(doc, "Phase 13: Tests")
    code(doc, """
uv add --dev pytest
uv run pytest
""")
    body(doc,
         "Use a synthetic inventory and a stub embedder so the whole suite runs "
         "offline with no API calls. Test the behaviour, not the model:")
    bullets(doc, [
        "Every stated requirement is honored.",
        '"good gas mileage" does not become fuel=gas.',
        "Relaxation drops the weakest constraint and protects fuel.",
        "Explicit filters override inferred ones.",
        "Budget phrasings, ranges and floors all bind.",
        "Mileage and bare model years are not read as budgets.",
        "Budget digits never reach the embedder.",
        "Every agent tool: casing, sink updates, clamped limits, unknown ids refused.",
    ])

    sub(doc, "Testing the UI itself")
    body(doc,
         "Booting the app and getting HTTP 200 proves almost nothing, Streamlit "
         "serves a shell and only executes the script when a browser connects. "
         "AppTest runs the real script headlessly and lets you click things:")
    code(doc, """
from streamlit.testing.v1 import AppTest

at = AppTest.from_file("app.py")
at.session_state["sink"] = {"results": sample_df, "label": "x", "meta": None}
at.run()

at.button[0].click().run()
assert len(at.get("dialog")) == 1          # modal opened
at.run()
assert len(at.get("dialog")) == 0          # and does not reopen after dismissal
""")
    note(doc,
         "This is how the reopening bug and the wrong-card-on-reorder bug were caught "
         "before they ever reached the browser. Element accessors are named for the "
         "element type, at.image, at.metric, at.button, and getting the name wrong "
         "returns an empty list rather than an error, which reads exactly like a "
         "missing element. Confirm against the element tree before believing a zero.")

    # ----------------------------------------------------------------- phase 13
    phase(doc, "Phase 14: Final structure")
    code(doc, """
car-recom-system/
├── PROJECT_GUIDE.docx
├── build_guide_docx.py         regenerates the guide above
├── README.md
├── app.py                      Streamlit entry point
├── pyproject.toml / uv.lock
├── showroom/
│   ├── config.py               paths + model ids, anchored to the repo root
│   ├── query_parser.py         lifts hard requirements out of free text
│   ├── search_engine.py        filtering, relaxation, ranking
│   └── sales_agent.py          LangGraph agent + tools
├── pipeline/
│   ├── build_dataset.py        Kaggle raw -> cleaned, sampled table
│   ├── vibe_tagging.py         zero-shot vibe scores
│   └── build_embeddings.py     OpenAI embeddings
├── notebooks/
│   └── data-exploration.ipynb
├── tests/
└── data/                       generated, gitignored
""")
    body(doc, "Full rebuild from nothing:")
    code(doc, """
uv sync
echo "OPENAI_API_KEY=sk-..." > .env

uv run python -m pipeline.build_dataset      # ~2 min
uv run python -m pipeline.vibe_tagging       # ~110 min, checkpointed
uv run python -m pipeline.build_embeddings   # ~3 min, ~$0.05

uv run pytest
uv run streamlit run app.py
""")
    note(doc,
         "The built index is about 120 MB, over GitHub's 100 MB per-file limit. Either "
         "use Git LFS or gitignore data/ and rely on the pipeline scripts.")

    sub(doc, "What is still imperfect")
    bullets(doc, [
        "Seller mislabels, the source has a Camry tagged electric and a Tesla tagged hybrid. Filtering is only as accurate as the labels.",
        "Manufacturer is not a hard constraint, \"show me a Toyota\" ranks Toyotas highly but does not filter to them.",
        "Vibe scores are softmax shares, not independent confidences, so they compare poorly across rows.",
        "No images, the source photos expired. See Phase 11.",
    ])

    sub(doc, "The thread running through all of it")
    body(doc,
         "Nearly every bug in this project was the same bug wearing a different "
         "costume: the system presenting something as matching the request when it "
         "did not.")
    bullets(doc, [
        "Ranking a manual car highly for \"I want automatic\", because nothing enforced the requirement.",
        "Returning $7,500 cars for \"under 5000\", because price was never parsed.",
        "Relaxing a constraint and describing the results as though it still held.",
        "Quoting a seller's typo as an established spec.",
        "Showing a stock photo of a different car than the one for sale.",
    ])
    body(doc,
         "The fixes all share a shape too: enforce what was actually asked, and when "
         "you cannot deliver it, say so plainly instead of returning the nearest thing "
         "and hoping it passes. A recommender that admits \"we have no electric "
         "sedans, here are the hybrids\" is more useful than one that silently "
         "substitutes, and much easier to trust on the answers it does give.")

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    build()
